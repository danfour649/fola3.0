#!/usr/bin/env python3
"""
Reusable markdown to DOCX/PDF exporter for project documents.
"""

from __future__ import annotations

from dataclasses import dataclass
from html import escape
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


@dataclass(frozen=True)
class Block:
    block_type: str
    text: str


def parse_markdown(md_text: str) -> list[Block]:
    """
    Parse a small markdown subset into typed blocks:
    - h1, h2, h3
    - bullet
    - number
    - text
    """
    blocks: list[Block] = []
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            blocks.append(Block("spacer", ""))
            continue
        if line.startswith("# "):
            blocks.append(Block("h1", line[2:].strip()))
            continue
        if line.startswith("## "):
            blocks.append(Block("h2", line[3:].strip()))
            continue
        if line.startswith("### "):
            blocks.append(Block("h3", line[4:].strip()))
            continue
        if line.startswith("- "):
            blocks.append(Block("bullet", line[2:].strip()))
            continue
        if re.match(r"^\d+\.\s+", line):
            blocks.append(Block("number", line))
            continue

        blocks.append(Block("text", line))

    return blocks


def write_docx(blocks: list[Block], out_path: Path) -> None:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for block in blocks:
        if block.block_type == "spacer":
            document.add_paragraph("")
        elif block.block_type == "h1":
            document.add_heading(block.text, level=1)
        elif block.block_type == "h2":
            document.add_heading(block.text, level=2)
        elif block.block_type == "h3":
            document.add_heading(block.text, level=3)
        elif block.block_type == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.block_type == "number":
            number_text = re.sub(r"^\d+\.\s+", "", block.text).strip()
            document.add_paragraph(number_text, style="List Number")
        else:
            document.add_paragraph(block.text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def write_pdf(blocks: list[Block], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    style_sheet = getSampleStyleSheet()
    heading_1 = ParagraphStyle(
        "H1Custom",
        parent=style_sheet["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=8,
        textColor=colors.HexColor("#1F2937"),
    )
    heading_2 = ParagraphStyle(
        "H2Custom",
        parent=style_sheet["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.HexColor("#111827"),
    )
    heading_3 = ParagraphStyle(
        "H3Custom",
        parent=style_sheet["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        spaceBefore=6,
        spaceAfter=2,
        textColor=colors.HexColor("#111827"),
    )
    body_style = ParagraphStyle(
        "BodyCustom",
        parent=style_sheet["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=body_style,
        leftIndent=14,
    )

    story = []
    pending_bullets: list[ListItem] = []

    def flush_bullets() -> None:
        nonlocal pending_bullets
        if not pending_bullets:
            return
        story.append(
            ListFlowable(
                pending_bullets,
                bulletType="bullet",
                start="circle",
                leftIndent=12,
            )
        )
        story.append(Spacer(1, 4))
        pending_bullets = []

    for block in blocks:
        if block.block_type != "bullet":
            flush_bullets()

        text = escape(block.text)
        if block.block_type == "spacer":
            story.append(Spacer(1, 6))
        elif block.block_type == "h1":
            story.append(Paragraph(text, heading_1))
        elif block.block_type == "h2":
            story.append(Paragraph(text, heading_2))
        elif block.block_type == "h3":
            story.append(Paragraph(text, heading_3))
        elif block.block_type == "bullet":
            pending_bullets.append(ListItem(Paragraph(text, bullet_style)))
        else:
            story.append(Paragraph(text, body_style))

    flush_bullets()
    document.build(story)


def parse_formats(value: str) -> list[str]:
    raw_formats = [fmt.strip().lower() for fmt in value.split(",") if fmt.strip()]
    unique_formats: list[str] = []
    for fmt in raw_formats:
        if fmt not in {"docx", "pdf"}:
            raise ValueError(f"Unsupported format '{fmt}'. Use docx, pdf, or both.")
        if fmt not in unique_formats:
            unique_formats.append(fmt)
    if not unique_formats:
        raise ValueError("No output formats specified.")
    return unique_formats


def convert_markdown_file(
    input_path: Path,
    output_dir: Path,
    output_basename: str | None = None,
    formats: list[str] | None = None,
) -> dict[str, Path]:
    if not input_path.exists():
        raise FileNotFoundError(f"Input markdown file not found: {input_path}")

    selected_formats = formats or ["docx", "pdf"]
    md_text = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    base_name = output_basename or input_path.stem
    outputs: dict[str, Path] = {}

    if "docx" in selected_formats:
        docx_path = output_dir / f"{base_name}.docx"
        write_docx(blocks, docx_path)
        outputs["docx"] = docx_path

    if "pdf" in selected_formats:
        pdf_path = output_dir / f"{base_name}.pdf"
        write_pdf(blocks, pdf_path)
        outputs["pdf"] = pdf_path

    return outputs
