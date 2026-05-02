#!/usr/bin/env python3
"""
Generate DOCX and PDF versions of the business plan from markdown.

Outputs:
- deliverables/business-plan-folas-cleaning-company.docx
- deliverables/business-plan-folas-cleaning-company.pdf
"""

from __future__ import annotations

from pathlib import Path
import re
import textwrap

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, ListFlowable, ListItem


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = PROJECT_ROOT / "docs" / "06-business-plan-folas-cleaning-company.md"
OUT_DOCX = PROJECT_ROOT / "deliverables" / "business-plan-folas-cleaning-company.docx"
OUT_PDF = PROJECT_ROOT / "deliverables" / "business-plan-folas-cleaning-company.pdf"


def parse_markdown(md_text: str) -> list[tuple[str, str]]:
    """
    Parse a small markdown subset into typed blocks:
    - h1, h2, h3
    - bullet
    - text
    """
    blocks: list[tuple[str, str]] = []
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line or line.strip() == "---":
            blocks.append(("spacer", ""))
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            blocks.append(("text", line))
        else:
            blocks.append(("text", line))
    return blocks


def write_docx(blocks: list[tuple[str, str]], out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for block_type, text in blocks:
        if block_type == "spacer":
            doc.add_paragraph("")
        elif block_type == "h1":
            doc.add_heading(text, level=1)
        elif block_type == "h2":
            doc.add_heading(text, level=2)
        elif block_type == "h3":
            doc.add_heading(text, level=3)
        elif block_type == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def write_pdf(blocks: list[tuple[str, str]], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1Custom",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=8,
        textColor=colors.HexColor("#1f2937"),
    )
    h2 = ParagraphStyle(
        "H2Custom",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.HexColor("#111827"),
    )
    h3 = ParagraphStyle(
        "H3Custom",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        spaceBefore=6,
        spaceAfter=2,
        textColor=colors.HexColor("#111827"),
    )
    body = ParagraphStyle(
        "BodyCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=body,
        leftIndent=14,
    )

    story = []
    pending_bullets: list[ListItem] = []

    def flush_bullets() -> None:
        nonlocal pending_bullets
        if pending_bullets:
            story.append(
                ListFlowable(
                    pending_bullets,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=12,
                )
            )
            story.append(Spacer(1, 4))
            pending_bullets = []

    for block_type, text in blocks:
        clean_text = textwrap.fill(text, width=160) if text else text
        if block_type != "bullet":
            flush_bullets()

        if block_type == "spacer":
            story.append(Spacer(1, 6))
        elif block_type == "h1":
            story.append(Paragraph(clean_text, h1))
        elif block_type == "h2":
            story.append(Paragraph(clean_text, h2))
        elif block_type == "h3":
            story.append(Paragraph(clean_text, h3))
        elif block_type == "bullet":
            pending_bullets.append(ListItem(Paragraph(clean_text, bullet_style)))
        else:
            escaped = (
                clean_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(escaped, body))

    flush_bullets()
    doc.build(story)


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    md_text = SOURCE_MD.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    write_docx(blocks, OUT_DOCX)
    write_pdf(blocks, OUT_PDF)

    print(f"Generated DOCX: {OUT_DOCX}")
    print(f"Generated PDF:  {OUT_PDF}")


if __name__ == "__main__":
    main()
